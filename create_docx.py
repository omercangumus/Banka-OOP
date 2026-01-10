from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title Page
title = doc.add_heading('NovaBank - Dijital Bankacılık Uygulaması', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Proje Raporu')
run.bold = True
run.font.size = Pt(24)

doc.add_paragraph()
doc.add_paragraph()

# Info table
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('Ders:', '.NET Uygulama Geliştirme (NTP)'),
    ('Üniversite:', 'Fırat Üniversitesi'),
    ('Bölüm:', 'Bilgisayar Mühendisliği'),
    ('Öğrenci:', 'Ömer Can Gümüş'),
    ('Tarih:', 'Ocak 2026')
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# Table of Contents
doc.add_heading('İÇİNDEKİLER', level=1)
toc_items = [
    '1. Proje Özeti',
    '2. Giriş ve Problem Tanımı',
    '3. Sistem Gereksinimleri',
    '4. Teknoloji Stack',
    '5. Mimari Tasarım',
    '6. Veritabanı Tasarımı',
    '7. UML Diyagramları',
    '8. Modül Açıklamaları',
    '9. Test Stratejisi',
    '10. Güvenlik Önlemleri',
    '11. Sonuç ve Değerlendirme',
    '12. Kaynaklar'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. Project Summary
doc.add_heading('1. PROJE ÖZETİ', level=1)
doc.add_paragraph('NovaBank, modern dijital bankacılık deneyimi sunan, Windows Forms ve DevExpress tabanlı kurumsal düzeyde bir masaüstü bankacılık uygulamasıdır. Proje, Clean Architecture prensiplerine uygun olarak 4 katmanlı mimari yapıda geliştirilmiştir.')

doc.add_heading('Temel Özellikler:', level=2)
features = [
    'Kullanıcı kayıt ve kimlik doğrulama',
    'Email doğrulama sistemi (OTP)',
    'Hesap yönetimi (TL/USD/EUR)',
    'Para transferi (IBAN ile)',
    'Yatırım platformu (Hisse, Kripto, Emtia)',
    'Teknik analiz araçları (RSI, MACD, Bollinger Bands)',
    'AI Finansal Asistan',
    'Admin paneli',
    'Kredi kartı yönetimi',
    'Kredi başvuru sistemi',
    'PDF raporlama'
]
for feature in features:
    doc.add_paragraph(f'✅ {feature}', style='List Bullet')

doc.add_heading('Proje İstatistikleri:', level=2)
stats_table = doc.add_table(rows=8, cols=2)
stats_table.style = 'Table Grid'
stats_data = [
    ('Metrik', 'Değer'),
    ('Toplam C# Dosyası', '100+'),
    ('Satır Kod', '~25,000'),
    ('Form Sayısı', '25+'),
    ('User Control', '18'),
    ('Servis Sayısı', '30+'),
    ('Entity Sayısı', '15'),
    ('Unit Test', '15+')
]
for i, (col1, col2) in enumerate(stats_data):
    stats_table.rows[i].cells[0].text = col1
    stats_table.rows[i].cells[1].text = col2
    if i == 0:
        stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        stats_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# 2. Introduction
doc.add_heading('2. GİRİŞ VE PROBLEM TANIMI', level=1)

doc.add_heading('2.1 Problem Tanımı', level=2)
doc.add_paragraph('Günümüzde bankacılık işlemleri büyük ölçüde dijitalleşmiştir. Ancak geleneksel bankacılık sistemleri karmaşık arayüzler sunar, yatırım araçlarına sınırlı erişim sağlar, teknik analiz olanakları kısıtlıdır ve AI destekli danışmanlık hizmeti yoktur.')

doc.add_heading('2.2 Çözüm Önerisi', level=2)
solutions = [
    'Modern UI/UX: DevExpress bileşenleri ile profesyonel görünüm',
    'Entegre Yatırım: Tek platformda hisse, kripto ve emtia işlemleri',
    'Gelişmiş Analiz: Candlestick grafikleri, teknik indikatörler',
    'AI Destekli: OpenRouter API ile akıllı finansal danışman',
    'Güvenlik: SHA256 şifreleme, email doğrulama, audit logging'
]
for solution in solutions:
    doc.add_paragraph(solution, style='List Bullet')

doc.add_page_break()

# 3. System Requirements
doc.add_heading('3. SİSTEM GEREKSİNİMLERİ', level=1)

doc.add_heading('3.1 Fonksiyonel Gereksinimler', level=2)
req_table = doc.add_table(rows=11, cols=3)
req_table.style = 'Table Grid'
req_data = [
    ('ID', 'Gereksinim', 'Öncelik'),
    ('FR01', 'Kullanıcı kayıt olabilmeli', 'Yüksek'),
    ('FR02', 'Email doğrulama yapılmalı', 'Yüksek'),
    ('FR03', 'Hesaplar arası transfer', 'Yüksek'),
    ('FR04', 'Yatırım alım/satım', 'Orta'),
    ('FR05', 'Teknik analiz görüntüleme', 'Orta'),
    ('FR06', 'Kredi başvurusu', 'Orta'),
    ('FR07', 'Kredi kartı oluşturma', 'Orta'),
    ('FR08', 'Admin kullanıcı yönetimi', 'Yüksek'),
    ('FR09', 'PDF rapor oluşturma', 'Düşük'),
    ('FR10', 'AI asistan sohbet', 'Düşük')
]
for i, row_data in enumerate(req_data):
    for j, cell_data in enumerate(row_data):
        req_table.rows[i].cells[j].text = cell_data
        if i == 0:
            req_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# 4. Technology Stack
doc.add_heading('4. TEKNOLOJİ STACK', level=1)

doc.add_heading('4.1 Geliştirme Ortamı', level=2)
tech_table = doc.add_table(rows=11, cols=3)
tech_table.style = 'Table Grid'
tech_data = [
    ('Kategori', 'Teknoloji', 'Versiyon'),
    ('Framework', '.NET', '8.0'),
    ('Dil', 'C#', '12'),
    ('IDE', 'Visual Studio', '2022'),
    ('UI Framework', 'DevExpress WinForms', '25.2'),
    ('Database', 'PostgreSQL', '16'),
    ('ORM', 'Dapper', '2.1.35'),
    ('Email', 'MailKit', '4.x'),
    ('PDF', 'DevExpress XtraReports', '25.2'),
    ('Test', 'xUnit', '2.9'),
    ('CI/CD', 'GitLab CI', '-')
]
for i, row_data in enumerate(tech_data):
    for j, cell_data in enumerate(row_data):
        tech_table.rows[i].cells[j].text = cell_data
        if i == 0:
            tech_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_heading('4.2 Harici API Entegrasyonları', level=2)
api_table = doc.add_table(rows=4, cols=3)
api_table.style = 'Table Grid'
api_data = [
    ('API', 'Kullanım Amacı', 'Endpoint'),
    ('OpenRouter', 'AI Finansal Danışman', 'openrouter.ai'),
    ('Finnhub', 'Hisse Senedi Verileri', 'finnhub.io'),
    ('Binance', 'Kripto Para Verileri', 'binance.com')
]
for i, row_data in enumerate(api_data):
    for j, cell_data in enumerate(row_data):
        api_table.rows[i].cells[j].text = cell_data
        if i == 0:
            api_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# 5. Architecture
doc.add_heading('5. MİMARİ TASARIM', level=1)

doc.add_heading('5.1 Clean Architecture', level=2)
doc.add_paragraph('NovaBank projesi Clean Architecture prensiplerine uygun olarak tasarlanmıştır:')

arch_text = '''
┌─────────────────────────────────────────────────────────────┐
│                    PRESENTATION LAYER                        │
│                      (BankApp.UI)                            │
│     Forms, Controls, User Interface Components               │
├─────────────────────────────────────────────────────────────┤
│                  APPLICATION / INFRASTRUCTURE                │
│                 (BankApp.Infrastructure)                     │
│       Services, Repositories, External APIs, Email           │
├─────────────────────────────────────────────────────────────┤
│                      DOMAIN LAYER                            │
│                     (BankApp.Core)                           │
│           Entities, Interfaces, Business Rules               │
└─────────────────────────────────────────────────────────────┘
'''
doc.add_paragraph(arch_text, style='No Spacing')

doc.add_heading('5.2 Proje Yapısı', level=2)
structure = [
    'BankaBenim/',
    '├── src/',
    '│   ├── BankApp.Core/           # Domain Layer (15 Entity)',
    '│   ├── BankApp.Infrastructure/ # Services & Repositories',
    '│   ├── BankApp.UI/             # 25+ Windows Forms',
    '│   └── BankApp.Tests/          # Unit Tests',
    '├── docs/                       # Dokümantasyon',
    '└── *.md                        # Proje dökümanları'
]
for line in structure:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

doc.add_page_break()

# 6. Database Design
doc.add_heading('6. VERİTABANI TASARIMI', level=1)

doc.add_heading('6.1 Veritabanı Bilgileri', level=2)
db_table = doc.add_table(rows=4, cols=2)
db_table.style = 'Table Grid'
db_data = [
    ('RDBMS', 'PostgreSQL 16'),
    ('Veritabanı Adı', 'NovaBankDb'),
    ('Karakter Seti', 'UTF-8'),
    ('Connection', 'Dapper (Micro-ORM)')
]
for i, (label, value) in enumerate(db_data):
    db_table.rows[i].cells[0].text = label
    db_table.rows[i].cells[1].text = value

doc.add_heading('6.2 Tablolar', level=2)
tables_info = [
    ('Users', 'Kullanıcı bilgileri, kimlik doğrulama'),
    ('Customers', 'Müşteri profili, KYC bilgileri'),
    ('Accounts', 'Banka hesapları (TL/USD/EUR)'),
    ('Transactions', 'Para transferi kayıtları'),
    ('Loans', 'Kredi başvuruları'),
    ('CreditCards', 'Kredi kartı bilgileri'),
    ('CustomerPortfolio', 'Yatırım portföyü'),
    ('AuditLogs', 'Denetim kayıtları')
]
for table_name, desc in tables_info:
    p = doc.add_paragraph()
    run = p.add_run(f'• {table_name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# 7. UML Diagrams
doc.add_heading('7. UML DİYAGRAMLARI', level=1)
doc.add_paragraph('Aşağıdaki UML diyagramları projenin farklı yönlerini görselleştirmektedir. Diyagram görsel dosyaları UML_Diagrams klasöründe bulunmaktadır.')

uml_list = [
    ('7.1 Sınıf Diyagramı (Class Diagram)', 'Tüm entity sınıfları ve aralarındaki ilişkiler'),
    ('7.2 Use Case Diyagramı', 'Müşteri ve Admin kullanım senaryoları'),
    ('7.3 Sequence Diyagramı', 'Para transferi işlem akışı'),
    ('7.4 Activity Diyagramı', 'Kredi başvurusu iş akışı'),
    ('7.5 ER Diyagramı', 'Veritabanı entity-relationship şeması'),
    ('7.6 Component Diyagramı', '3 katmanlı mimari yapı'),
    ('7.7 State Diyagramı', 'Kredi durumu geçişleri'),
    ('7.8 Deployment Diyagramı', 'Sistem dağıtım mimarisi')
]

for title, desc in uml_list:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    doc.add_paragraph('[Diyagram UML_Diagrams klasöründe bulunmaktadır]')

doc.add_page_break()

# 8. Module Descriptions
doc.add_heading('8. MODÜL AÇIKLAMALARI', level=1)

doc.add_heading('8.1 Kimlik Doğrulama Modülü (AuthService)', level=2)
auth_features = [
    'Kayıt: Yeni kullanıcı kaydı, şifre hash\'leme',
    'Giriş: Kullanıcı kimlik doğrulama',
    'Email Doğrulama: 6 haneli OTP kodu gönderimi',
    'Şifre Sıfırlama: Email ile şifre yenileme',
    'Audit Logging: Tüm giriş/çıkış kayıtları'
]
for feature in auth_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('8.2 Transfer Modülü (TransactionService)', level=2)
transfer_features = [
    'IBAN Transfer: IBAN numarası ile transfer',
    'Bakiye Kontrolü: Yetersiz bakiye kontrolü',
    'Anlık Güncelleme: Gerçek zamanlı bakiye güncelleme',
    'İşlem Geçmişi: Tüm transferlerin kaydı'
]
for feature in transfer_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('8.3 Yatırım Modülü (InvestmentService)', level=2)
invest_features = [
    'Hisse Alım/Satım: THYAO, GARAN gibi hisseler',
    'Kripto İşlemleri: BTC, ETH alım/satım',
    'Emtia: Altın, gümüş işlemleri',
    'Portföy Takibi: Anlık portföy değeri',
    'Teknik Analiz: RSI, MACD, Bollinger Bands'
]
for feature in invest_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('8.4 AI Asistan Modülü', level=2)
ai_features = [
    'Portföy Analizi: Yatırım portföyü değerlendirmesi',
    'Risk Analizi: Portföy risk seviyesi belirleme',
    'Piyasa Durumu: Güncel piyasa özeti',
    'Teknik Analiz Yorumu: Grafik pattern yorumlama',
    'Sohbet: Doğal dil finansal sorgular'
]
for feature in ai_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# 9. Test Strategy
doc.add_heading('9. TEST STRATEJİSİ', level=1)

doc.add_heading('9.1 Unit Testler', level=2)
doc.add_paragraph('xUnit, Moq ve FluentAssertions kullanılarak yazılmıştır:')
test_items = [
    'AuthServiceTests: Login, Register, Verify işlemleri',
    'TransactionServiceTests: Transfer, bakiye kontrol',
    'InvestmentDashboardV2Tests: RSI, MACD hesaplama'
]
for item in test_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.2 Smoke Testler', level=2)
smoke_table = doc.add_table(rows=10, cols=3)
smoke_table.style = 'Table Grid'
smoke_data = [
    ('Kategori', 'Test Sayısı', 'Durum'),
    ('Application Startup', '3', '✅ PASS'),
    ('Authentication', '2', '✅ PASS'),
    ('Admin Panel', '3', '✅ PASS'),
    ('Export', '2', '✅ PASS'),
    ('User Workflow', '2', '✅ PASS'),
    ('Integration', '2', '✅ PASS'),
    ('Performance', '2', '✅ PASS'),
    ('Security', '2', '✅ PASS'),
    ('TOPLAM', '20', '✅ ALL PASS')
]
for i, row_data in enumerate(smoke_data):
    for j, cell_data in enumerate(row_data):
        smoke_table.rows[i].cells[j].text = cell_data
        if i == 0 or i == 9:
            smoke_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# 10. Security
doc.add_heading('10. GÜVENLİK ÖNLEMLERİ', level=1)

doc.add_heading('10.1 Kimlik Doğrulama', level=2)
security_items = [
    'SHA256 şifre hashleme',
    'Email OTP doğrulama (6 haneli kod)',
    '15 dakika kod geçerliliği',
    'Hesap kilitleme yanlış giriş sonrası'
]
for item in security_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Veri Güvenliği', level=2)
data_security = [
    'Parameterized SQL - SQL Injection koruması',
    'Input Validation - Tüm form girişleri doğrulanır',
    'Audit Logging - Tüm kritik işlemler kaydedilir',
    'API Key Güvenliği - Anahtarlar gitignore\'da tutulur'
]
for item in data_security:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.3 Yetkilendirme', level=2)
role_table = doc.add_table(rows=3, cols=2)
role_table.style = 'Table Grid'
role_data = [
    ('Rol', 'Yetkiler'),
    ('Customer', 'Hesap görüntüleme, transfer, yatırım, kart'),
    ('Admin', 'Tüm yetkiler + kullanıcı yönetimi + kredi onay')
]
for i, row_data in enumerate(role_data):
    for j, cell_data in enumerate(row_data):
        role_table.rows[i].cells[j].text = cell_data
        if i == 0:
            role_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# 11. Conclusion
doc.add_heading('11. SONUÇ VE DEĞERLENDİRME', level=1)

doc.add_heading('11.1 Proje Kazanımları', level=2)
doc.add_paragraph('Bu proje sürecinde aşağıdaki teknik yetkinlikler kazanılmıştır:')
gains = [
    'Clean Architecture prensiplerine uygun katmanlı mimari tasarım',
    '.NET 8.0 ile modern C# geliştirme',
    'DevExpress WinForms ile profesyonel UI tasarım',
    'PostgreSQL veritabanı yönetimi',
    'Dapper Micro-ORM kullanımı',
    'API Entegrasyonu (Finnhub, Binance, OpenRouter)',
    'Unit Testing (xUnit, Moq, FluentAssertions)',
    'Git/GitLab versiyon kontrolü ve CI/CD'
]
for gain in gains:
    doc.add_paragraph(gain, style='List Bullet')

doc.add_heading('11.2 Gelecek Geliştirmeler', level=2)
future = [
    'Mobil uygulama (MAUI)',
    'REST API katmanı',
    'İki faktörlü kimlik doğrulama (2FA)',
    'Push notification sistemi',
    'Bütçe planlama modülü'
]
for item in future:
    doc.add_paragraph(f'☐ {item}', style='List Bullet')

doc.add_page_break()

# 12. References
doc.add_heading('12. KAYNAKLAR', level=1)

doc.add_heading('Resmi Dokümantasyon', level=2)
refs = [
    'Microsoft .NET Documentation: https://docs.microsoft.com/dotnet',
    'DevExpress Documentation: https://docs.devexpress.com',
    'PostgreSQL Documentation: https://www.postgresql.org/docs',
    'Dapper Documentation: https://dapperlib.github.io/Dapper'
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('API Referansları', level=2)
apis = [
    'Finnhub API: https://finnhub.io/docs/api',
    'Binance API: https://binance-docs.github.io/apidocs',
    'OpenRouter API: https://openrouter.ai/docs'
]
for api in apis:
    doc.add_paragraph(api, style='List Bullet')

doc.add_page_break()

# Appendix
doc.add_heading('EKLER', level=1)

doc.add_heading('Ek-1: Varsayılan Test Kullanıcıları', level=2)
user_table = doc.add_table(rows=5, cols=3)
user_table.style = 'Table Grid'
user_data = [
    ('Kullanıcı', 'Şifre', 'Rol'),
    ('admin', 'admin123', 'Admin'),
    ('test', 'test123', 'Customer'),
    ('demo', 'demo123', 'Customer'),
    ('staff', '123456', 'Staff')
]
for i, row_data in enumerate(user_data):
    for j, cell_data in enumerate(row_data):
        user_table.rows[i].cells[j].text = cell_data
        if i == 0:
            user_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Proje Teslim Tarihi: Ocak 2026')
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Hazırlayan: Ömer Can Gümüş')
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('Fırat Üniversitesi - Bilgisayar Mühendisliği')

# Save
output_path = r'C:\Users\omerc\Desktop\BankaBenim\NovaBank_Proje_Raporu.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
